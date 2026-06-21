#!/usr/bin/env python3
"""Generate the official team-leader project reference DOCX."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "scripts" / "sheet-data.json"
OUT_PATH = ROOT / "project-team-leaders-reference.docx"
DAMIETTA_LOGO = ROOT / "public" / "brand" / "damietta-university.jpg"
FCAI_LOGO = ROOT / "public" / "brand" / "fcai-logo.jpg"

NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(31, 78, 121)
GOLD = RGBColor(210, 160, 23)
LIGHT_BLUE = "D9EAF7"
PALE_GOLD = "FFF4CC"
LIGHT_GRAY = "F4F6F8"
BORDER = "9EADBD"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=11, bold=False, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rtl = rpr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rpr.append(rtl)
    rtl.set(qn("w:val"), "1")


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=6, line=1.15):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_rtl(p, align=align, before=0, after=0, line=1.05)
    run = p.add_run("" if text is None else str(text))
    set_run_font(run, size=size, bold=bold, color=color or RGBColor(0, 0, 0))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width_and_rtl(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi)
    bidi.set(qn("w:val"), "1")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    table._tbl.insert(1, grid)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_aragraph(doc, text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=6):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=align, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or RGBColor(0, 0, 0))
    return p


def add_header_block(doc, project_count):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1440, 6480, 1440]
    set_table_width_and_rtl(table, widths)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=0, bottom=0)
        for paragraph in cell.paragraphs:
            set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    logo_cells = table.rows[0].cells
    if FCAI_LOGO.exists():
        logo_cells[0].paragraphs[0].add_run().add_picture(str(FCAI_LOGO), width=Inches(0.83))
    if DAMIETTA_LOGO.exists():
        logo_cells[2].paragraphs[0].add_run().add_picture(str(DAMIETTA_LOGO), width=Inches(0.83))

    center = logo_cells[1]
    center.text = ""
    for text, size, bold, color in [
        ("جامعة دمياط", 13, True, NAVY),
        ("كلية الحاسبات والمعلومات والذكاء الاصطناعي", 12, True, BLUE),
        ("نظام مراجعة وتسليم مشاريع التخرج", 10, False, RGBColor(80, 80, 80)),
    ]:
        p = center.add_paragraph()
        set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)

    spacer = doc.add_paragraph()
    set_paragraph_rtl(spacer, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

    title = doc.add_paragraph()
    set_paragraph_rtl(title, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
    r = title.add_run("كشف قادة الفرق وبيانات التسجيل")
    set_run_font(r, size=18, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    set_paragraph_rtl(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
    r = subtitle.add_run(f"مشاريع التخرج - الترم الثاني 2026 | عدد المشاريع: {project_count}")
    set_run_font(r, size=10.5, bold=True, color=RGBColor(90, 90, 90))


def add_instruction_box(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width_and_rtl(table, [9360])
    set_table_borders(table, color="D6B656", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, top=130, bottom=130, start=170, end=170)
    cell.text = ""

    heading = cell.add_paragraph()
    set_paragraph_rtl(heading, before=0, after=4)
    run = heading.add_run("رسالة العميد لقادة الفرق")
    set_run_font(run, size=13, bold=True, color=NAVY)

    paragraphs = [
        "أبنائي الطلاب، يسعدنا إطلاق نظام مراجعة مشاريع التخرج بكلية الحاسبات والمعلومات والذكاء الاصطناعي بجامعة دمياط. مرفق بهذا البيان كشف رسمي يوضح رقم كل مشروع وقائد الفريق والرقم الجامعي المستخدم في التسجيل.",
        "على قائد الفريق فقط الدخول إلى https://dam-rs.vercel.app/register وإنشاء الحساب باستخدام الرقم الجامعي الموضح أمام مشروعه والرقم القومي الخاص به. الرقم القومي لا يُنشر في هذا الكشف حفاظاً على سرية بيانات الطلاب.",
        "بعد التسجيل، يقوم قائد الفريق بتسجيل الدخول واستكمال بيانات المشروع: الملخص، رابط فيديو العرض، بيانات أعضاء الفريق، ورفع مستند المشروع بصيغة PDF، والكود المصدري بصيغة ZIP، وملف العرض التقديمي.",
        "لا يُسمح لأي عضو آخر بإنشاء تسليم مستقل أو تكرار التسليم باسم المشروع. التسليم النهائي يتم مرة واحدة لكل مشروع من خلال قائد الفريق فقط.",
        "آخر موعد للتسليم: يُحدد لاحقاً.",
    ]
    for text in paragraphs:
        p = cell.add_paragraph()
        set_paragraph_rtl(p, before=0, after=3, line=1.15)
        r = p.add_run(text)
        set_run_font(r, size=9.4, color=RGBColor(30, 30, 30))


def add_projects_table(doc, projects):
    add_aragraph(doc, "كشف المشروعات وقادة الفرق", size=13, bold=True, color=NAVY, before=10, after=5)
    add_aragraph(
        doc,
        "يُستخدم الرقم الجامعي لقائد الفريق في إنشاء الحساب. كلمة المرور/التحقق هي الرقم القومي الخاص بقائد الفريق عند التسجيل.",
        size=9.5,
        color=RGBColor(75, 75, 75),
        after=6,
    )

    headers = ["م", "رقم المشروع", "عنوان المشروع", "قائد الفريق", "الرقم الجامعي للقائد"]
    widths = [500, 820, 4120, 2680, 1240]
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width_and_rtl(table, widths)
    set_table_borders(table, color=BORDER, size="5")
    set_repeat_table_header(table.rows[0])

    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in (0, 1, 4) else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_text(cell, label, size=9.1, bold=True, color=NAVY, align=align)

    for i, project in enumerate(projects, start=1):
        row = table.add_row()
        values = [
            i,
            project.get("number", ""),
            project.get("title_ar", ""),
            project.get("leader_full_name", ""),
            project.get("leader_university_id", ""),
        ]
        fill = "FFFFFF" if i % 2 else LIGHT_GRAY
        for col, value in enumerate(values):
            cell = row.cells[col]
            set_cell_shading(cell, fill)
            align = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 1, 4) else WD_ALIGN_PARAGRAPH.RIGHT
            size = 8.3 if col == 2 else 8.6
            set_cell_text(cell, value, size=size, align=align)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    r = p.add_run("كلية الحاسبات والمعلومات والذكاء الاصطناعي - جامعة دمياط | نظام مراجعة مشاريع التخرج")
    set_run_font(r, size=8.5, color=RGBColor(100, 100, 100))


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)


def main():
    data = json.loads(DATA_PATH.read_text(encoding="utf-8-sig"))
    projects = sorted(data["projects"], key=lambda p: int(p["number"]))

    doc = Document()
    configure_document(doc)
    add_header_block(doc, len(projects))
    add_projects_table(doc, projects)

    doc.core_properties.title = "كشف قادة الفرق وبيانات التسجيل"
    doc.core_properties.subject = "مشاريع التخرج - الترم الثاني 2026"
    doc.core_properties.author = "كلية الحاسبات والمعلومات والذكاء الاصطناعي - جامعة دمياط"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
