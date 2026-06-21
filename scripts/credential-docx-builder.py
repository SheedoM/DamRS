#!/usr/bin/env python3
"""Build credential DOCX sheets using the official Arabic reference style."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DAMIETTA_LOGO = ROOT / "public" / "brand" / "damietta-university.jpg"
FCAI_LOGO = ROOT / "public" / "brand" / "fcai-logo.jpg"

NAVY = RGBColor(23, 54, 93)
BLUE = RGBColor(31, 78, 121)
GOLD = RGBColor(210, 160, 23)
LIGHT_BLUE = "D9EAF7"
PALE_GOLD = "FFF4CC"
LIGHT_GRAY = "F4F6F8"
BORDER = "9EADBD"


def set_run_font(run, size=11, bold=False, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rtl = rpr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        rpr.append(rtl)
    rtl.set(qn("w:val"), "1")


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=6, line=1.15):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        ppr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, size=10, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_rtl(p, align=align, before=0, after=0, line=1.05)
    run = p.add_run("" if text is None else str(text))
    set_run_font(run, size=size, bold=bold, color=color or RGBColor(0, 0, 0))


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width_and_rtl(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi)
    bidi.set(qn("w:val"), "1")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    table._tbl.insert(1, grid)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_paragraph(doc, text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT, before=0, after=6):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, align=align, before=before, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color or RGBColor(0, 0, 0))
    return p


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)
    run = p.add_run("كلية الحاسبات والمعلومات والذكاء الاصطناعي - جامعة دمياط | نظام مراجعة مشاريع التخرج")
    set_run_font(run, size=8.5, color=RGBColor(100, 100, 100))


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(11)


def add_header_block(doc, title, subtitle):
    table = doc.add_table(rows=1, cols=3)
    widths = [1440, 6480, 1440]
    set_table_width_and_rtl(table, widths)
    for cell in table.rows[0].cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=0, bottom=0)
        for paragraph in cell.paragraphs:
            set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0)

    cells = table.rows[0].cells
    if FCAI_LOGO.exists():
        cells[0].paragraphs[0].add_run().add_picture(str(FCAI_LOGO), width=Inches(0.83))
    if DAMIETTA_LOGO.exists():
        cells[2].paragraphs[0].add_run().add_picture(str(DAMIETTA_LOGO), width=Inches(0.83))

    center = cells[1]
    center.text = ""
    for text, size, bold, color in [
        ("جامعة دمياط", 13, True, NAVY),
        ("كلية الحاسبات والمعلومات والذكاء الاصطناعي", 12, True, BLUE),
        ("نظام مراجعة وتسليم مشاريع التخرج", 10, False, RGBColor(80, 80, 80)),
    ]:
        p = center.add_paragraph()
        set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=2)
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)

    spacer = doc.add_paragraph()
    set_paragraph_rtl(spacer, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)

    title_p = doc.add_paragraph()
    set_paragraph_rtl(title_p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=18, bold=True, color=NAVY)

    subtitle_p = doc.add_paragraph()
    set_paragraph_rtl(subtitle_p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
    subtitle_run = subtitle_p.add_run(subtitle)
    set_run_font(subtitle_run, size=10.5, bold=True, color=RGBColor(90, 90, 90))


def add_note_box(doc, heading, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_width_and_rtl(table, [9360])
    set_table_borders(table, color="D6B656", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GOLD)
    set_cell_margins(cell, top=120, bottom=120, start=170, end=170)
    cell.text = ""

    p = cell.add_paragraph()
    set_paragraph_rtl(p, before=0, after=4)
    run = p.add_run(heading)
    set_run_font(run, size=12.5, bold=True, color=NAVY)

    for line in lines:
        p = cell.add_paragraph()
        set_paragraph_rtl(p, before=0, after=3, line=1.12)
        run = p.add_run(line)
        set_run_font(run, size=9.3, color=RGBColor(30, 30, 30))


def add_table(doc, title, description, headers, widths, rows, centered_columns):
    add_paragraph(doc, title, size=13, bold=True, color=NAVY, before=10, after=5)
    if description:
        add_paragraph(doc, description, size=9.5, color=RGBColor(75, 75, 75), after=6)

    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width_and_rtl(table, widths)
    set_table_borders(table, color=BORDER, size="5")
    set_repeat_table_header(table.rows[0])

    for idx, label in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_BLUE)
        align = WD_ALIGN_PARAGRAPH.CENTER if idx in centered_columns else WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_text(cell, label, size=9.1, bold=True, color=NAVY, align=align)

    for row_idx, values in enumerate(rows, start=1):
        row = table.add_row()
        fill = "FFFFFF" if row_idx % 2 else LIGHT_GRAY
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            set_cell_shading(cell, fill)
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx in centered_columns else WD_ALIGN_PARAGRAPH.RIGHT
            size = 8.2 if len(str(value or "")) > 40 else 8.7
            set_cell_text(cell, value, size=size, align=align)


def panel_type_label(value):
    return {
        "supervisor": "مشرف",
        "referee": "محكم",
        "committee_head": "رئيس اللجنة",
    }.get(value or "", value or "")


def leader_status_label(value):
    return {
        "claimed": "تم التسجيل",
        "pending": "بانتظار التسجيل",
        "generated-id": "رقم جامعي مولد",
        "matched": "مطابق",
    }.get(value or "", value or "")


def build_panel_doc(panel_rows, out_path):
    doc = Document()
    configure_document(doc)
    add_header_block(doc, "كشف بيانات دخول أعضاء اللجان", f"عدد الحسابات: {len(panel_rows)}")
    rows = [
        [idx, item.get("full_name", ""), item.get("email", ""), item.get("password") or "(غير محددة)", panel_type_label(item.get("type"))]
        for idx, item in enumerate(panel_rows, start=1)
    ]
    add_table(
        doc,
        "بيانات الدخول",
        "يستخدم عضو اللجنة البريد الإلكتروني وكلمة المرور التالية في صفحة دخول أعضاء اللجان.",
        ["م", "الاسم", "البريد الإلكتروني", "كلمة المرور", "النوع"],
        [500, 2700, 2780, 2050, 1330],
        rows,
        {0, 2, 3, 4},
    )
    doc.core_properties.title = "كشف بيانات دخول أعضاء اللجان"
    doc.core_properties.subject = "مشاريع التخرج - بيانات دخول أعضاء اللجان"
    doc.core_properties.author = "كلية الحاسبات والمعلومات والذكاء الاصطناعي - جامعة دمياط"
    doc.save(out_path)


def build_leaders_doc(leader_rows, out_path):
    doc = Document()
    configure_document(doc)
    add_header_block(doc, "كشف قادة الفرق وبيانات التسجيل", f"عدد المشاريع: {len(leader_rows)}")
    rows = [
        [
            idx,
            item.get("project_number", ""),
            item.get("title", ""),
            item.get("leader_name", ""),
            item.get("leader_university_id", ""),
        ]
        for idx, item in enumerate(leader_rows, start=1)
    ]
    add_table(
        doc,
        "كشف المشروعات وقادة الفرق",
        "يُستخدم الرقم الجامعي لقائد الفريق في إنشاء الحساب. كلمة المرور/التحقق هي الرقم القومي الخاص بقائد الفريق عند التسجيل.",
        ["م", "رقم المشروع", "عنوان المشروع", "قائد الفريق", "الرقم الجامعي للقائد"],
        [500, 820, 4120, 2680, 1240],
        rows,
        {0, 1, 4},
    )
    doc.core_properties.title = "كشف قادة الفرق وبيانات التسجيل"
    doc.core_properties.subject = "مشاريع التخرج - بيانات تسجيل قادة الفرق"
    doc.core_properties.author = "كلية الحاسبات والمعلومات والذكاء الاصطناعي - جامعة دمياط"
    doc.save(out_path)


def load_input(path):
    if path == "-":
        return json.load(sys.stdin)
    with open(path, encoding="utf-8-sig") as handle:
        return json.load(handle)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True)
    parser.add_argument("--panel-out", required=True)
    parser.add_argument("--leaders-out", required=True)
    args = parser.parse_args()

    payload = load_input(args.input)
    build_panel_doc(payload.get("panel", []), args.panel_out)
    build_leaders_doc(payload.get("leaders", []), args.leaders_out)
    print(args.panel_out)
    print(args.leaders_out)


if __name__ == "__main__":
    main()
